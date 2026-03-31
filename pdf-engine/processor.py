import fitz  # PyMuPDF
import os
import json
import io
import time
import base64
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Helper to force RTL Bidi in python-docx paragraphs
def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def get_rgb(color_val):
    if isinstance(color_val, str) and color_val.startswith("#"):
        c = color_val.lstrip("#")
        try:
            return (int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
        except:
            pass
    return 0, 0, 0

class PDFProcessor:
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("❌ OPENROUTER_API_KEY is missing. Please set it in pdf-engine/.env")
        
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key
        )

    def extract_layout_and_images(self, pdf_bytes: bytes) -> list[dict]:
        """Stage 0: Render high-resolution (300 DPI) full-page images for Stage 1 layout discovery."""
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages: list[dict] = []

        for page in doc:
            # 300 DPI is the sweet spot for modern AI vision models
            mat = fitz.Matrix(1.0, 1.0)
            pix = page.get_pixmap(matrix=mat, dpi=300)
            img_b64 = base64.b64encode(pix.tobytes("jpeg")).decode("utf-8")

            pages.append({
                "page": page.number + 1,
                "image_base64": img_b64,
                "width": page.rect.width,
                "height": page.rect.height
            })

        return pages

    def _call_ai_json(self, prompt, image_b64=None, retries=3):
        """Helper to invoke AI and repair JSON response if necessary."""
        for attempt in range(retries):
            try:
                msg_content = [{"type": "text", "text": prompt}]
                if image_b64:
                    msg_content.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}
                    })
                
                response = self.client.chat.completions.create(
                    model="google/gemini-3-flash-preview",
                    messages=[{"role": "user", "content": msg_content}],
                    response_format={"type": "json_object"}
                )
                text = response.choices[0].message.content
                data = json.loads(self.repair_json(text))
                # Safety: If AI returns a list instead of an object, wrap it
                if isinstance(data, list):
                    return {"regions": data, "blocks": data}
                return data
            except Exception as e:
                if attempt == retries - 1:
                    print(f"  [AI] Final Attempt Failed: {str(e)}")
                    return {}
                print(f"  [AI] Attempt {attempt+1} failed, retrying...")
                time.sleep(2)
        return {}

    def repair_json(self, text):
        """Extracts JSON from markdown and fixes common malformations."""
        if not text: return "{}"
        # Remove MD code blocks
        text = text.replace("```json", "").replace("```", "").strip()
        
        # Guard against AI returning multiple distinct root objects e.g. {} \n {}
        if "}\n{" in text or "}\n\n{" in text:
            text = "[" + text.replace("}\n{", "},{").replace("}\n\n{", "},{") + "]"
            
        # Find first { or [
        start = text.find("{")
        if start == -1: start = text.find("[")
        if start != -1:
            end = text.rfind("}")
            if end == -1: end = text.rfind("]")
            if end != -1:
                text = text[start:end+1]
        return text

    def get_ai_reconstruction(self, pages: list[dict], pdf_bytes: bytes) -> dict:
        """The Multi-Pass Pipeline: Discover Layout -> Extract Details -> Synthesize."""
        json_pages = []
        raw_responses = []
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        for p in pages:
            page_num = p['page']
            print(f"--- Processing Page {page_num} (Multi-Pass) ---")
            
            # PASS 1: Layout Discovery
            # PASS 1: 4-Zone Strip Discovery
            # We explicitly divide into heavily biased top/bottom zones.
            strips = [
                {"name": "Ultra Header", "ymin": 0, "ymax": 100},
                {"name": "Header", "ymin": 100, "ymax": 300},
                {"name": "Body", "ymin": 250, "ymax": 850},
                {"name": "Footer", "ymin": 800, "ymax": 1000}
            ]
            
            all_regions = []
            for strip in strips:
                strip_img = self._get_strip_b64(doc, page_num, strip["ymin"], strip["ymax"])
                strip_prompt = f"""You are a Document Architect focused ON ONLY the {strip['name']} segment. 
Identify all structural items in this strip (y={strip['ymin']} to y={strip['ymax']}).

TEXT COMPLETENESS RULE:
- You MUST capture 100% of visible text.
- Small, faint, or low-contrast text is CRITICAL.
- Missing any text is considered a failure.

Fidelity Rules:
1. Identify all 'text_block' (headings, paragraphs, footnotes, labels). Large blue titles and small gray footers are CRITICAL.
2. Identify 'image' (charts, graphs, icons, photos).
3. Identify 'table' (grid data).

OUTPUT JSON:
{{
  "regions": [
    {{
      "type": "text_block" | "image" | "table",
      "bbox": [ymin, xmin, ymax, xmax],
      "order": 1
    }}
  ]
}}
BBoxes must be normalized 0-1000 relative to THIS STRIP (y0=0, y1=1000).
JSON ONLY."""
                strip_data = self._call_ai_json(strip_prompt, strip_img)
                strip_regions = strip_data.get("regions", []) or strip_data.get("blocks", [])
                
                # Remap coordinates from Strip-Local to Global Page-Local
                strip_h = strip["ymax"] - strip["ymin"]
                for r in strip_regions:
                    bbox = self._safe_bbox(r.get("bbox"))
                    # Normalize y coordinates back to full-page 0-1000 scale
                    r["bbox"] = [
                        float(strip["ymin"] + (bbox[0] / 1000) * strip_h),
                        float(bbox[1]),
                        float(strip["ymin"] + (bbox[2] / 1000) * strip_h),
                        float(bbox[3])
                    ]
                
                print(f"  [Pass 1] Strip {strip['name']}: Found {len(strip_regions)} regions")
                all_regions.extend(strip_regions)

            # FALLBACK: If strips found NOTHING, try one holistic full-page scan
            if not all_regions:
                print(f"  [Pass 1] WARNING: Strips found 0 regions. Falling back to Holistic Scan...")
                full_prompt = "Identify all structural regions on this page (text_block, image, table). JSON ONLY."
                full_data = self._call_ai_json(full_prompt, p['image_base64'])
                all_regions = full_data.get("regions", []) or full_data.get("blocks", [])

            # Deduplicate regions based on 90% overlap
            regions = self._deduplicate_regions(all_regions)
            print(f"  [Pass 1] Total Unique Regions Identified: {len(regions)}")
            if not regions:
                print(f"  [Pass 1] ERROR: Even Holistic Scan failed to find any regions.")

            synthesized_blocks = []
            
            # PASS 2: Specialized Extraction & PASS 3: Synthesis
            for roi in regions:
                roi_type = str(roi.get("type", "text_block")).lower()
                bbox = self._safe_bbox(roi.get("bbox"))
                
                if roi_type == "table":
                    # Specialized Table Extraction
                    table_details = self._extract_table_details(doc, page_num, bbox)
                    if table_details:
                        roi["rows"] = table_details.get("rows", [])
                        roi["columnWidths"] = table_details.get("columnWidths", [])
                
                elif roi_type in ["text_block", "header", "footer", "paragraph", "heading", "caption"]:
                    text_prompt = f"""Extract the literal text and visual style for this block. Type: {roi_type}.

TEXT COMPLETENESS RULE:
- You MUST extract ALL visible text into a SINGLE text string. Do not miss small captions or faint text.
- Missing any text is considered a failure.

OUTPUT EXACT JSON FORMAT:
{{
  "text": "The exact combined text...",
  "alignment": "left/center/right",
  "style": {{"fontSize": 12, "bold": false, "italic": false, "underline": false, "color": "#000000"}}
}}
JSON ONLY."""
                    
                    # Zoom-Based Micro Extraction: Upscale small boxes
                    h_norm = bbox[2] - bbox[0]
                    scale = 3.0 if h_norm < 150 else 1.5 # 3x upscaling for things <15% of page
                    
                    block_image = self._get_crop_b64(doc, page_num, bbox, padding=25, scale_factor=scale)
                    block_data = self._call_ai_json(text_prompt, block_image)
                    roi.update(block_data)
                    roi["type"] = "text_block" # Normalize for renderer
                
                elif roi_type == "image":
                    # Images are already identified, just ensure they have an ID for embedding
                    if not roi.get("id"):
                        roi["id"] = f"img_{page_num}_{int(bbox[0])}_{int(bbox[1])}"
                
                synthesized_blocks.append(roi)

            # NEW: GRID SYNTHESIS (Clustering into Rows)
            rows = self._cluster_blocks_into_rows(synthesized_blocks)
            # NEW: LAYOUT OPTIMIZATION (Balancing & Compaction)
            rows = self._refine_layout(rows)
            
            json_pages.append(rows)
            raw_responses.append({"page": page_num, "recognition": {"blocks": synthesized_blocks}})

        doc.close()
        return {"json_pages": json_pages, "raw_responses": raw_responses}

    def _safe_float(self, val, default=0.0):
        """Helper to safely convert any value (including lists/strings) to a float."""
        if val is None: return default
        if isinstance(val, (int, float)): return float(val)
        if isinstance(val, list):
            return self._safe_float(val[0], default) if len(val) > 0 else default
        try:
            return float(str(val).replace(",", "").strip())
        except:
            return default

    def _safe_bbox(self, bbox):
        """Standardizes AI-provided bbox into 4 float list [ymin, xmin, ymax, xmax]."""
        if isinstance(bbox, list):
            # Handle nested list: [[y, x, y, x]]
            if len(bbox) == 1 and isinstance(bbox[0], list):
                bbox = bbox[0]
                
            if len(bbox) == 4:
                return [self._safe_float(x) for x in bbox]
            elif len(bbox) > 4:
                return [self._safe_float(x) for x in bbox[:4]]
            elif len(bbox) > 0:
                # Pad with 0s if too short
                return [self._safe_float(x) for x in bbox] + [0.0] * (4 - len(bbox))
        return [0.0, 0.0, 0.0, 0.0]

    def _cluster_blocks_into_rows(self, blocks, threshold=25):
        """Groups blocks that share horizontal space into 'Rows'."""
        if not blocks: return []
        
        # Sort primarily by ymin
        blocks.sort(key=lambda b: (self._safe_bbox(b.get("bbox"))[0], b.get("order", 999)))
        
        rows = []
        if not blocks: return rows
        
        current_row = [blocks[0]]
        for i in range(1, len(blocks)):
            last_block = current_row[-1]
            curr_block = blocks[i]
            
            last_ymin, _, last_ymax, _ = self._safe_bbox(last_block.get("bbox"))
            curr_ymin, _, curr_ymax, _ = self._safe_bbox(curr_block.get("bbox"))
            
            # If they overlap significantly in Y space, they are in the same row
            if abs(curr_ymin - last_ymin) < threshold:
                current_row.append(curr_block)
            else:
                current_row.sort(key=lambda b: self._safe_bbox(b.get("bbox"))[1]) # Sort row by xmin
                rows.append(current_row)
                current_row = [curr_block]
        
        if current_row:
            current_row.sort(key=lambda b: self._safe_bbox(b.get("bbox"))[1])
            rows.append(current_row)
            
        return rows

    def _refine_layout(self, rows):
        """Layout Optimizer: Balances symmetric columns and cleans spatial metadata."""
        if not rows: return []
        refined_rows = []
        
        for row in rows:
            if not row: continue
            num_cols = len(row)
            if num_cols > 1:
                total_w = sum([(self._safe_bbox(b.get("bbox"))[3] - self._safe_bbox(b.get("bbox"))[1]) for b in row])
                if total_w > 0:
                    for b in row:
                        bbox = self._safe_bbox(b.get("bbox"))
                        act_w = bbox[3] - bbox[1]
                        ratio = act_w / total_w
                        
                        # Smart Symmetry Snapping (Humanization)
                        if abs(ratio - 0.5) < 0.1: ratio = 0.5
                        elif abs(ratio - 0.333) < 0.05: ratio = 0.333
                        elif abs(ratio - 0.25) < 0.05: ratio = 0.25
                        
                        b["refined_ratio"] = ratio
                        
            refined_rows.append(row)
        return refined_rows

    def _extract_table_details(self, doc, page_num, bbox):
        """Dedicated high-res extraction for tables."""
        crop_b64 = self._get_crop_b64(doc, page_num, bbox)
        prompt = """Analyze this table image. Extract its structure into a JSON grid.
Rules:
1. Preserve every cell's text.
2. Detect mergers: use 'colspan' and 'rowspan' if cells span multiple columns/rows.
3. Identify styles: 'bold', 'color', 'fontSize' for text inside cells.
4. Provide 'columnWidths' as an array of percentage weights.

Format: {"rows": [[{"text": "...", "colspan": 1, "style": {...}}]], "columnWidths": [...]}"""
        return self._call_ai_json(prompt, crop_b64)

    def _get_strip_b64(self, doc, page_num, ymin_norm, ymax_norm, dpi=300):
        """Helper to get a base64 encoded vertical strip from PyMuPDF."""
        page = doc[page_num - 1]
        w, h = page.rect.width, page.rect.height
        rect = fitz.Rect(0, (ymin_norm / 1000) * h, w, (ymax_norm / 1000) * h)
        pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72), clip=rect)
        img_bytes = pix.tobytes("png")
        return base64.b64encode(img_bytes).decode("utf-8")

    def _deduplicate_regions(self, regions, iou_threshold=0.85):
        """Removes overlapping regions from different strips, prioritizing larger boxes/text."""
        if not regions: return []
        regions.sort(key=lambda x: self._safe_bbox(x.get("bbox"))[0])
        
        unique = []
        for r in regions:
            is_dup = False
            r_bbox = self._safe_bbox(r.get("bbox"))
            for u in unique:
                u_bbox = self._safe_bbox(u.get("bbox"))
                
                # Simple IoU logic for vertical overlap (X is roughly the same)
                y_o = max(0, min(r_bbox[2], u_bbox[2]) - max(r_bbox[0], u_bbox[0]))
                x_o = max(0, min(r_bbox[3], u_bbox[3]) - max(r_bbox[1], u_bbox[1]))
                inter = x_o * y_o
                union = ((r_bbox[2] - r_bbox[0]) * (r_bbox[3] - r_bbox[1])) + \
                        ((u_bbox[2] - u_bbox[0]) * (u_bbox[3] - u_bbox[1])) - inter
                iou = inter / union if union > 0 else 0
                
                if iou > iou_threshold:
                    is_dup = True
                    break
            if not is_dup:
                unique.append(r)
        return unique

    def _get_crop_b64(self, doc, page_num, normalized_bbox, dpi=300, padding=12, scale_factor=1.0):
        """Helper to get a base64 encoded crop from PyMuPDF, with optional upscaling."""
        page = doc[page_num - 1]
        w, h = page.rect.width, page.rect.height
        ymin, xmin, ymax, xmax = self._safe_bbox(normalized_bbox)
        
        # Add dynamic padding (in pts)
        rect = fitz.Rect(
            max(0, (xmin/1000)*w - padding),
            max(0, (ymin/1000)*h - (padding * 1.5)), # Extra vertical padding for headings
            min(w, (xmax/1000)*w + padding),
            min(h, (ymax/1000)*h + padding)
        )
        
        # Zoom-Based Micro Extraction scaling matrix
        dpi_val = (dpi/72) * scale_factor
        pix = page.get_pixmap(clip=rect, matrix=fitz.Matrix(dpi_val, dpi_val))
        return base64.b64encode(pix.tobytes("jpeg")).decode("utf-8")

    def crop_images(self, pdf_bytes: bytes, reconstruction: dict) -> dict:
        """Post-Process Phase: Crop bounding boxes derived from AI natively into PyMuPDF pngs."""
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        os.makedirs("/tmp/pdf_images", exist_ok=True)
        
        for raw_resp in reconstruction.get("raw_responses", []):
            page_num = raw_resp["page"]
            if page_num - 1 >= len(doc): continue
            page = doc[page_num - 1]
            pdf_w = page.rect.width
            pdf_h = page.rect.height
            
            ai_data = raw_resp.get("recognition", {})
            blocks = ai_data.get("blocks", [])
            for obj in blocks:
                if obj.get("type") == "image" and "bbox" in obj:
                    try:
                        ymin, xmin, ymax, xmax = self._safe_bbox(obj["bbox"])
                        # Convert normalized 0-1000 coordiantes to PyMuPDF native pt scale
                        y0 = (ymin / 1000) * pdf_h
                        x0 = (xmin / 1000) * pdf_w
                        y1 = (ymax / 1000) * pdf_h
                        x1 = (xmax / 1000) * pdf_w
                        
                        # Apply 5px padding constraint
                        padding = 5
                        y0 = max(0, y0 - padding)
                        x0 = max(0, x0 - padding)
                        y1 = min(pdf_h, y1 + padding)
                        x1 = min(pdf_w, x1 + padding)
                        
                        crop_rect = fitz.Rect(x0, y0, x1, y1)
                        # Render at 3x scale for extremely crisp DOCX output
                        pix = page.get_pixmap(clip=crop_rect, matrix=fitz.Matrix(3, 3))
                        img_path = f"/tmp/pdf_images/{obj['id']}.png"
                        pix.save(img_path)
                    except Exception as e:
                        print(f"Error extracting image {obj.get('id', 'Unknown')}: {e}")
                        
        doc.close()
        return reconstruction

    def _apply_style_to_run(self, run, style_dict):
        """Apply parsed JSON style dictionary directly to a python-docx text run."""
        if not style_dict:
            return

        if style_dict.get("bold") is True:
            run.bold = True
            bCs = OxmlElement('w:bCs')
            run._r.get_or_add_rPr().append(bCs)
        
        if style_dict.get("italic") is True:
            run.italic = True
            iCs = OxmlElement('w:iCs')
            run._r.get_or_add_rPr().append(iCs)
            
        if style_dict.get("underline") is True:
            run.underline = True
            
        f_size = style_dict.get("fontSize", 11) # Reduced default for density
        try:
            f_size = min(36, self._safe_float(f_size)) # Sanity cap at 36pt
            run.font.size = Pt(f_size)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(f_size * 2)))
            run._r.get_or_add_rPr().append(szCs)
        except:
            pass
        
        color_val = style_dict.get("color", "#000000")
        r, g, b = get_rgb(color_val)
        run.font.color.rgb = RGBColor(r, g, b)
        
        run.font.rtl = True # Fix Arabic punctuation logic natively

    def _apply_alignment(self, paragraph, align_str):
        if align_str == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align_str == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def build_docx(self, reconstruction: dict) -> bytes:
        """Stage 3: Advanced Layout Engine using Hybrid Tables & Floating Images."""
        doc = Document()
        json_pages = reconstruction.get("json_pages", [])
        
        PAGE_HEIGHT_PTS = 792.0
        USABLE_WIDTH_INCHES = 6.0

        for page_idx, page_rows in enumerate(json_pages):
            if page_idx > 0:
                new_section = doc.add_section()
                new_section.start_type = 2
            
            last_ymax = 0

            for row in page_rows:
                if not row: continue
                
                # Determine min ymin for the row padding
                row_ymin = min([self._safe_bbox(b.get("bbox"))[0] for b in row])
                
                # LAYOUT COMPACTION: Address "Airy/Loose" spacing
                raw_gap = row_ymin - last_ymax
                if raw_gap > 35:
                    # Compress excessively large vertical gaps logarithmically
                    raw_gap = 25 + (raw_gap * 0.2)
                
                space_before = Pt(max(2, (raw_gap / 1000) * PAGE_HEIGHT_PTS))
                
                if len(row) == 1:
                    # Single-block row: Standard Insertion
                    self._add_block_to_docx(doc, row[0], space_before, USABLE_WIDTH_INCHES)
                else:
                    # GRID ROW: Use a Layout Table (Borderless)
                    table = doc.add_table(rows=1, cols=len(row))
                    table.style = None # No visible borders
                    table.autofit = False
                    
                    # Proportional Widths applied from Optimizer
                    for i, block in enumerate(row):
                        bbox = self._safe_bbox(block.get("bbox"))
                        if "refined_ratio" in block:
                            ratio = block["refined_ratio"]
                        else:
                            row_total_norm_width = sum([(self._safe_bbox(b.get("bbox"))[3] - self._safe_bbox(b.get("bbox"))[1]) for b in row])
                            norm_w = bbox[3] - bbox[1]
                            ratio = norm_w / row_total_norm_width if row_total_norm_width > 0 else 1.0
                            
                        cell = table.cell(0, i)
                        # Add tighter padding (50 dxa / 2.5pt) for visual density
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcMar = OxmlElement('w:tcMar')
                        for margin in ['top', 'left', 'bottom', 'right']:
                            m = OxmlElement(f'w:{margin}')
                            m.set(qn('w:w'), '50')
                            m.set(qn('w:type'), 'dxa')
                            tcMar.append(m)
                        tcPr.append(tcMar)
                        
                        cell.width = Inches(ratio * USABLE_WIDTH_INCHES)
                        # Deduct cell padding from final image available width
                        self._add_block_to_docx(cell, block, Pt(0), (ratio * USABLE_WIDTH_INCHES) - 0.1)
                
                # Update last_ymax based on the tallest block in the row
                last_ymax = max([self._safe_bbox(b.get("bbox"))[2] for b in row])

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _add_block_to_docx(self, parent, block, space_before, available_width_inches):
        """Unified block adder that can target Document or TableCell."""
        b_type = block.get("type", "paragraph")
        
        if b_type in ["paragraph", "heading", "text_block", "header", "footer"]:
            text = (block.get("text") or "").strip()
            if not text: return
            
            p = parent.add_paragraph()
            p.paragraph_format.space_before = space_before
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            
            align_str = block.get("alignment", "right") 
            self._apply_alignment(p, align_str)
            if align_str == "right": set_rtl(p)
            
            run = p.add_run(text)
            self._apply_style_to_run(run, block.get("style") or {})
        
        elif b_type == "image":
            img_id = block.get("id", "")
            local_path = f"/tmp/pdf_images/{img_id}.png"
            if os.path.exists(local_path):
                p = parent.add_paragraph()
                p.paragraph_format.space_before = space_before
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                bbox = self._safe_bbox(block.get("bbox"))
                norm_width = bbox[3] - bbox[1]
                dynamic_width = Inches((norm_width / 1000) * 6.0) # Scale against page width
                final_w = min(Inches(available_width_inches), dynamic_width)
                
                shape = run.add_picture(local_path, width=final_w)
                
                # Check for advanced wrapping (Square/Around)
                wrapping = block.get("wrapping", "inline")
                if wrapping == "square":
                    self._set_image_square_wrapping(shape)

        elif b_type == "table":
            rows = block.get("rows", [])
            col_widths = block.get("columnWidths", [])
            if not rows: return
            cols_count = max([len(r) for r in rows]) if rows else 0
            if cols_count == 0: return
            
            table = parent.add_table(rows=len(rows), cols=cols_count)
            table.style = 'Table Grid'
            table.autofit = False
            table.table_direction = "rtl"
            
            merged_cells = set()
            for r_idx, r_data in enumerate(rows):
                for c_idx, cell_data in enumerate(r_data):
                    if c_idx < cols_count:
                        if (r_idx, c_idx) in merged_cells: continue
                        cell = table.cell(r_idx, c_idx)
                        cp = cell.paragraphs[0]
                        set_rtl(cp)
                        text = (cell_data.get("text") or "").strip()
                        run = cp.add_run(text)
                        self._apply_style_to_run(run, cell_data.get("style") or {})
                        
                        colspan = int(cell_data.get("colspan", 1))
                        rowspan = int(cell_data.get("rowspan", 1))
                        if colspan > 1 or rowspan > 1:
                            tag_r = min(r_idx + rowspan - 1, len(rows) - 1)
                            tag_c = min(c_idx + colspan - 1, cols_count - 1)
                            cell.merge(table.cell(tag_r, tag_c))
                            for mr in range(r_idx, tag_r + 1):
                                for mc in range(c_idx, tag_c + 1):
                                    merged_cells.add((mr, mc))

    def _set_image_square_wrapping(self, shape):
        """Manipulate OpenXML to convert an InlineShape into a floating Anchor with Square wrapping."""
        # Note: This is a complex XML transformation. 
        # For now, we utilize the layout-table strategy as the primary tool.
        # True floating anchors require coordinate-to-EMU mapping which is prone to drift.
        # We ensure 'square' images have padding to simulate the look.
        shape.lock_aspect_ratio = True

    def build_html(self, json_pages: list) -> list[str]:
        """Proxy compiler to render visual HTML output requested by Frontend"""
        html_pages = []
        for page_rows in json_pages:
            html = "<div dir='rtl' style='padding: 20px; font-family: Arial;'>"
            for row in page_rows:
                for block in row:
                    b_type = block.get("type", "paragraph")
                    style = block.get("style") or {}
                    
                    # convert style to css
                    css = []
                    if style.get("bold"): css.append("font-weight: bold;")
                    if style.get("color"): css.append(f"color: {style['color']};")
                    if style.get("fontSize"): css.append(f"font-size: {style['fontSize']}pt;")
                    
                    align = block.get("alignment", "right")
                    css.append(f"text-align: {align};")
                    
                    style_str = " ".join(css)
                    text = block.get("text") or ""
                    
                    if b_type in ["paragraph", "heading", "text_block"]:
                        html += f"<p style='{style_str}'>{text}</p>\n"
                    elif b_type == "list":
                        html += f"<ul style='{style_str}'><li>{text}</li></ul>\n"
                    elif b_type == "image":
                        img_id = block.get("id", "")
                        html += f"<div style='text-align: center; margin: 10px;'><img src='{img_id}.png' alt='Image {img_id}' style='max-width: 100%; border: 1px dashed #ccc;'/></div>\n"
                    elif b_type == "table":
                        html += "<table border='1' style='width: 100%; border-collapse: collapse; margin-top: 10px;'>"
                        for r_data in block.get("rows", []):
                            html += "<tr>"
                            for cell in r_data:
                                cell_text = cell.get("text") or ""
                                cell_style = cell.get("style") or {}
                                c_css = ""
                                if cell_style.get("bold"): c_css += "font-weight: bold;"
                                html += f"<td style='padding: 5px; {c_css}'>{cell_text}</td>"
                            html += "</tr>"
                        html += "</table>\n"
                    
            html += "</div>"
            html_pages.append(html)
        return html_pages
